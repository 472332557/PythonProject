"""
文档读取模块 - 支持PDF、Word、TXT等常见格式文档的解析和关键信息提取
"""

from flask import Blueprint, request, jsonify
from functools import wraps
import logging
import os
import re
from datetime import datetime

logger = logging.getLogger(__name__)

document_bp = Blueprint('document', __name__)


def error_handler(f):
    """错误处理装饰器"""
    @wraps(f)
    def wrapper(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except ValueError as e:
            return jsonify({
                'code': 400,
                'message': f'参数错误: {str(e)}',
                'data': None
            }), 400
        except FileNotFoundError as e:
            return jsonify({
                'code': 404,
                'message': f'文件未找到: {str(e)}',
                'data': None
            }), 404
        except Exception as e:
            logger.error(f"文档处理错误: {str(e)}")
            return jsonify({
                'code': 500,
                'message': f'服务器错误: {str(e)}',
                'data': None
            }), 500
    return wrapper


def validate_file_path(file_path: str) -> str:
    """验证文件路径"""
    if not file_path:
        raise ValueError("文件路径不能为空")

    # 检查文件是否存在
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 检查文件大小 (限制50MB)
    file_size = os.path.getsize(file_path)
    if file_size > 50 * 1024 * 1024:
        raise ValueError("文件大小超过50MB限制")

    return file_path


def get_file_extension(file_path: str) -> str:
    """获取文件扩展名"""
    _, ext = os.path.splitext(file_path)
    return ext.lower().replace('.', '')


def read_document(file_path: str, max_length: int = 50000) -> dict:
    """
    读取并解析文档 (function-calling核心函数)

    Args:
        file_path: 文档路径
        max_length: 最大读取长度

    Returns:
        包含文档内容和关键信息的字典
    """
    file_path = validate_file_path(file_path)
    ext = get_file_extension(file_path)

    if ext == 'txt':
        return _read_txt(file_path, max_length)
    elif ext == 'pdf':
        return _read_pdf(file_path, max_length)
    elif ext in ['doc', 'docx']:
        return _read_word(file_path, max_length)
    elif ext == 'md':
        return _read_markdown(file_path, max_length)
    else:
        raise ValueError(f"不支持的文档格式: {ext}")


def _read_txt(file_path: str, max_length: int) -> dict:
    """读取TXT文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    return _process_document_content(file_path, content, max_length)


def _read_pdf(file_path: str, max_length: int) -> dict:
    """读取PDF文件"""
    try:
        import PyPDF2
    except ImportError:
        logger.warning("PyPDF2未安装，使用文本提取")
        return _extract_pdf_as_text(file_path, max_length)

    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)

            full_text = []
            for page in reader.pages[:min(num_pages, 100)]:  # 最多100页
                text = page.extract_text()
                if text:
                    full_text.append(text)

            content = '\n'.join(full_text)
            return _process_document_content(file_path, content, max_length)

    except Exception as e:
        logger.error(f"PDF读取失败: {str(e)}")
        return _extract_pdf_as_text(file_path, max_length)


def _extract_pdf_as_text(file_path: str, max_length: int) -> dict:
    """使用pdfminer提取PDF文本"""
    try:
        from pdfminer.high_level import extract_text
        content = extract_text(file_path)
        return _process_document_content(file_path, content, max_length)
    except ImportError:
        return {
            'error': 'PDF读取失败，请安装PyPDF2或pdfminer',
            'file_name': os.path.basename(file_path)
        }


def _read_word(file_path: str, max_length: int) -> dict:
    """读取Word文件"""
    try:
        from docx import Document
    except ImportError:
        return {
            'error': 'Word读取失败，请安装python-docx',
            'file_name': os.path.basename(file_path)
        }

    doc = Document(file_path)

    # 提取段落
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 提取表格
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    content = '\n'.join(paragraphs)

    result = _process_document_content(file_path, content, max_length)
    result['tables'] = tables
    result['paragraph_count'] = len(paragraphs)

    return result


def _read_markdown(file_path: str, max_length: int) -> dict:
    """读取Markdown文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    result = _process_document_content(file_path, content, max_length)
    result['format'] = 'markdown'

    # 提取标题结构
    headings = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
    result['structure'] = [
        {'level': len(h[0]), 'title': h[1]} for h in headings
    ]

    return result


def _process_document_content(file_path: str, content: str, max_length: int) -> dict:
    """处理文档内容，提取关键信息"""
    # 截断内容
    if len(content) > max_length:
        truncated = True
        content = content[:max_length] + '...'
    else:
        truncated = False

    # 提取关键信息
    file_size = os.path.getsize(file_path)

    result = {
        'file_name': os.path.basename(file_path),
        'file_path': file_path,
        'file_size': file_size,
        'format': get_file_extension(file_path),
        'content': content,
        'char_count': len(content),
        'word_count': len(content.split()),
        'line_count': content.count('\n') + 1,
        'truncated': truncated,
        'read_time': f"{len(content.split()) // 200}分钟",  # 粗略估算阅读时间
        'key_info': _extract_key_info(content)
    }

    return result


def _extract_key_info(content: str) -> dict:
    """从内容中提取关键信息"""
    key_info = {
        'emails': [],
        'phone_numbers': [],
        'urls': [],
        'numbers': []
    }

    # 提取邮箱
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    key_info['emails'] = re.findall(email_pattern, content)

    # 提取电话号码
    phone_pattern = r'\b1[3-9]\d{9}\b'
    key_info['phone_numbers'] = re.findall(phone_pattern, content)

    # 提取URL
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    key_info['urls'] = re.findall(url_pattern, content)

    # 提取数字
    number_pattern = r'\b\d+\.?\d*\b'
    numbers = re.findall(number_pattern, content)
    key_info['numbers'] = [n for n in numbers if len(n) < 15][:20]  # 限制数量

    # 清理空列表
    key_info = {k: v for k, v in key_info.items() if v}

    return key_info


@document_bp.route('/read', methods=['POST'])
@error_handler
def read_document_api():
    """
    读取文档接口

    Request Body:
        {
            "file_path": "/path/to/document.pdf",
            "max_length": 50000  // 可选，默认50000
        }

    Response:
        {
            "code": 200,
            "message": "success",
            "data": {
                "file_name": "document.pdf",
                "content": "文档内容...",
                "char_count": 12345,
                "key_info": {...},
                ...
            }
        }
    """
    data = request.get_json()
    if not data or 'file_path' not in data:
        return jsonify({
            'code': 400,
            'message': '缺少file_path参数',
            'data': None
        }), 400

    file_path = data['file_path']
    max_length = data.get('max_length', 50000)

    logger.info(f"读取文档: {file_path}")

    result = read_document(file_path, max_length)

    if 'error' in result:
        return jsonify({
            'code': 400,
            'message': result['error'],
            'data': None
        }), 400

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': result
    })


@document_bp.route('/info', methods=['POST'])
@error_handler
def get_document_info():
    """
    获取文档基本信息

    Request Body:
        {
            "file_path": "/path/to/document.pdf"
        }
    """
    data = request.get_json()
    if not data or 'file_path' not in data:
        return jsonify({
            'code': 400,
            'message': '缺少file_path参数',
            'data': None
        }), 400

    file_path = data['file_path']

    try:
        file_path = validate_file_path(file_path)
        ext = get_file_extension(file_path)

        return jsonify({
            'code': 200,
            'message': 'success',
            'data': {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'format': ext,
                'last_modified': datetime.fromtimestamp(
                    os.path.getmtime(file_path)
                ).strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except (ValueError, FileNotFoundError) as e:
        return jsonify({
            'code': 400 if isinstance(e, ValueError) else 404,
            'message': str(e),
            'data': None
        })


@document_bp.route('/search', methods=['POST'])
@error_handler
def search_in_document():
    """
    在文档中搜索关键词

    Request Body:
        {
            "file_path": "/path/to/document.pdf",
            "keyword": "关键词",
            "case_sensitive": false
        }
    """
    data = request.get_json()
    if not data or 'file_path' not in data or 'keyword' not in data:
        return jsonify({
            'code': 400,
            'message': '缺少必要参数',
            'data': None
        }), 400

    file_path = data['file_path']
    keyword = data['keyword']
    case_sensitive = data.get('case_sensitive', False)

    # 读取文档
    doc_data = read_document(file_path)
    if 'error' in doc_data:
        return jsonify({
            'code': 400,
            'message': doc_data['error'],
            'data': None
        }), 400

    # 搜索关键词
    content = doc_data['content']
    if not case_sensitive:
        content = content.lower()
        keyword = keyword.lower()

    # 找出匹配位置
    matches = []
    start = 0
    while True:
        pos = content.find(keyword, start)
        if pos == -1:
            break
        matches.append({
            'position': pos,
            'context': content[max(0, pos-50):min(len(content), pos+len(keyword)+50)]
        })
        start = pos + 1

        if len(matches) >= 100:  # 限制匹配数量
            break

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'file_name': os.path.basename(file_path),
            'keyword': keyword,
            'total_matches': len(matches),
            'matches': matches
        }
    })


@document_bp.route('/supported_formats', methods=['GET'])
def get_supported_formats():
    """获取支持的文档格式列表"""
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'formats': [
                {'extension': 'txt', 'name': '文本文件', 'support': '完整支持'},
                {'extension': 'pdf', 'name': 'PDF文档', 'support': '完整支持'},
                {'extension': 'doc', 'name': 'Word 97-2003', 'support': '完整支持'},
                {'extension': 'docx', 'name': 'Word 2007+', 'support': '完整支持'},
                {'extension': 'md', 'name': 'Markdown', 'support': '完整支持'}
            ],
            'max_file_size': '50MB'
        }
    })


# Function Calling 函数定义
FUNCTION_DEFINITIONS = {
    'name': 'read_document',
    'description': '读取并解析常见格式文档(PDF、Word、TXT、Markdown)，提取文档内容和关键信息',
    'parameters': {
        'type': 'object',
        'properties': {
            'file_path': {
                'type': 'string',
                'description': '文档的完整路径',
                'example': '/path/to/document.pdf'
            },
            'max_length': {
                'type': 'integer',
                'description': '最大读取字符数，默认50000',
                'example': 50000
            }
        },
        'required': ['file_path']
    }
}
